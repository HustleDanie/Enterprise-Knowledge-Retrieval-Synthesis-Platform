import logging
import time
import os
import json
import asyncio
from datetime import datetime
from fastapi import APIRouter, HTTPException, Depends, Header, UploadFile, File, BackgroundTasks
from typing import Optional, List
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from .schemas import QueryRequest, QueryResponse, HealthResponse

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/v1", tags=["search"])

# Initialize services (lazy loading)
embedding_service = None
vector_store = None
hybrid_retriever = None
llm_client = None

# Storage
uploaded_documents = {}
document_chunks = []
UPLOAD_DIR = "data/uploads"
METADATA_FILE = "data/uploads/metadata.json"
QUERY_CACHE = {}  # Simple in-memory cache for queries
CACHE_EXPIRY_MINUTES = 60

# Ensure upload directory exists
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ============================================================
# PRE-LOADED SAMPLE DOCUMENTS (Available immediately on startup)
# ============================================================

FOOTBALL_HISTORY_DOCUMENT = """
The History of Football: From Ancient Times to Modern Glory

ORIGINS AND EARLY FORMS (Ancient Times - 1800s)

Football's origins trace back thousands of years to ancient civilizations. The Chinese played a game called Cuju during the Han Dynasty around 206 BC, which involved kicking a leather ball through a small opening. The Greeks had Episkyros, while the Romans played Harpastum. In medieval England, "mob football" was played between neighboring towns with few rules and unlimited players.

THE BIRTH OF MODERN FOOTBALL (1863)

Modern football was born on October 26, 1863, when the Football Association (FA) was founded at the Freemason's Tavern in London, England. This historic meeting established the first standardized rules of the game, separating association football (soccer) from rugby football. The original rules included 13 laws of the game.

THE FIRST INTERNATIONAL MATCH (1872)

The first official international football match took place on November 30, 1872, between England and Scotland at Hamilton Crescent in Glasgow. The match ended 0-0. This historic game established international football competition and led to the formation of national teams worldwide.

THE FA CUP - OLDEST COMPETITION (1871)

The FA Cup, established in 1871, is the oldest football competition in the world. Wanderers FC won the first FA Cup final, defeating Royal Engineers 1-0 at Kennington Oval in London. The trophy remains one of the most prestigious in English football.

FIFA FORMATION (1904)

The Federation Internationale de Football Association (FIFA) was founded on May 21, 1904, in Paris, France. The founding members were Belgium, Denmark, France, Netherlands, Spain, Sweden, and Switzerland. FIFA became the global governing body of football, establishing international rules and organizing world competitions.

THE FIRST WORLD CUP (1930)

The inaugural FIFA World Cup was held in Uruguay in 1930. Uruguay, the host nation, won the tournament, defeating Argentina 4-2 in the final at the Estadio Centenario in Montevideo. Only 13 teams participated, with most European nations declining due to the long sea journey.

LEGENDARY PLAYERS THROUGH HISTORY

Pele (Brazil, 1940-2022): Widely considered the greatest footballer of all time. Won three World Cups (1958, 1962, 1970). Scored over 1,000 career goals. Named FIFA Player of the Century.

Diego Maradona (Argentina, 1960-2020): Led Argentina to World Cup victory in 1986. Famous for both the "Hand of God" goal and the "Goal of the Century" against England in the same match.

Johan Cruyff (Netherlands, 1947-2016): Pioneer of "Total Football." Won three consecutive Ballon d'Or awards (1971, 1973, 1974). Revolutionized tactical approaches to the game.

Lionel Messi (Argentina, born 1987): Won record 8 Ballon d'Or awards. Led Argentina to World Cup victory in 2022. Considered one of the greatest players in history alongside Pele and Maradona.

Cristiano Ronaldo (Portugal, born 1985): All-time top scorer in international football. Won five Ballon d'Or awards and multiple Champions League titles with Manchester United and Real Madrid.

MAJOR FOOTBALL COMPETITIONS

FIFA World Cup: Held every four years since 1930. Brazil holds the record with 5 titles (1958, 1962, 1970, 1994, 2002). Germany and Italy have won 4 each.

UEFA Champions League: Europe's premier club competition, evolved from the European Cup founded in 1955. Real Madrid holds the record with 14 titles.

Copa America: South America's continental championship, first held in 1916. Argentina and Uruguay have won the most titles.

UEFA European Championship (Euros): Europe's national team competition held every four years. Germany and Spain have won 3 titles each.

English Premier League: Founded in 1992, it's the most-watched football league globally. Manchester United has won 13 titles, followed by Manchester City and Chelsea.

FOOTBALL TODAY

Modern football is a global phenomenon with an estimated 4 billion fans worldwide. The sport generates over $30 billion annually. Technology like VAR (Video Assistant Referee) was introduced in 2018 to assist match officials. Women's football has grown significantly, with the FIFA Women's World Cup attracting over 1 billion viewers in 2019.

The 2022 FIFA World Cup in Qatar was the first held in the Middle East and the first in November-December to avoid extreme heat. Argentina won, with Lionel Messi cementing his legacy as one of the greatest players ever.
"""

SAMPLE_DOCUMENT = """
System Architecture Overview: Ingest pipeline (docs/PDFs -> chunking/cleaning) -> 
embedding generation -> vector DB ingestion -> user query -> retrieval (hybrid search) -> RAG 
chain (context + LLM) -> response + citations -> FastAPI backend -> frontend users; scheduled 
re-indexing + monitoring loop.

AI/ML Components: Full ML pipeline for chunk classification (XGBoost/Naive Bayes 
baseline); feature engineering on metadata; evaluation (recall@K, MRR).

GenAI / LLM / Agentic Components: OpenAI/GPT-4o or Llama-3 via HF; embeddings (text-embedding-ada-002 or BGE); advanced RAG (hybrid keyword+semantic, query rewriting, reranking).

MLOps & Engineering Components: MLflow tracking/experiments; Docker containerization; 
CI/CD (GitHub Actions); FastAPI async serving; AWS/Azure deployment (App Service/AKS or 
SageMaker endpoint); Prometheus/Grafana monitoring; model/data drift detection (Evidently 
AI); logging (ELK stack).

Tech Stack: Python, Langchain/LlamaIndex, Pinecone/Chroma (vector DB), PostgreSQL 
(metadata SQL), FastAPI, Docker/K8s, AWS Bedrock or Azure OpenAI + Azure ML, MLflow.

Advanced Features: Role-based access control (RBAC) on chunks; cost/latency optimization 
(caching, batching); hallucination mitigation (self-check agents); production constraints (privacy 
via PII redaction, scalability for millions of docs).
"""

def load_sample_documents():
    global document_chunks, uploaded_documents
    
    # === FOOTBALL HISTORY DOCUMENT ===
    football_doc_id = "football_history_001"
    text = FOOTBALL_HISTORY_DOCUMENT.strip()
    sentences = text.replace('\n', ' ').split('. ')
    chunks = []
    current_chunk = []
    current_size = 0
    chunk_size = 512
    
    for sentence in sentences:
        sentence = sentence.strip() + '.'
        if current_size + len(sentence) > chunk_size and current_chunk:
            chunks.append(' '.join(current_chunk))
            current_chunk = [sentence]
            current_size = len(sentence)
        else:
            current_chunk.append(sentence)
            current_size += len(sentence)
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    for i, chunk_text in enumerate(chunks):
        document_chunks.append({
            "chunk_id": f"{football_doc_id}_chunk_{i}",
            "doc_id": football_doc_id,
            "filename": "History_of_Football.txt",
            "content": chunk_text,
            "index": i
        })
    
    uploaded_documents[football_doc_id] = {
        "id": football_doc_id,
        "filename": "History_of_Football.txt",
        "size": len(FOOTBALL_HISTORY_DOCUMENT),
        "uploaded_at": datetime.now().isoformat(),
        "status": "ready",
        "chunk_count": len(chunks),
        "content": FOOTBALL_HISTORY_DOCUMENT
    }
    
    logger.info(f"Football history document loaded: {len(chunks)} chunks")
    
    # === RAG ARCHITECTURE DOCUMENT ===
    rag_doc_id = "sample_doc_001"
    text2 = SAMPLE_DOCUMENT.strip()
    sentences2 = text2.replace('\n', ' ').split('. ')
    chunks2 = []
    current_chunk2 = []
    current_size2 = 0
    
    for sentence in sentences2:
        sentence = sentence.strip() + '.'
        if current_size2 + len(sentence) > chunk_size and current_chunk2:
            chunks2.append(' '.join(current_chunk2))
            current_chunk2 = [sentence]
            current_size2 = len(sentence)
        else:
            current_chunk2.append(sentence)
            current_size2 += len(sentence)
    if current_chunk2:
        chunks2.append(' '.join(current_chunk2))
    
    for i, chunk_text in enumerate(chunks2):
        document_chunks.append({
            "chunk_id": f"{rag_doc_id}_chunk_{i}",
            "doc_id": rag_doc_id,
            "filename": "RAG_System_Architecture.txt",
            "content": chunk_text,
            "index": i
        })
    
    uploaded_documents[rag_doc_id] = {
        "id": rag_doc_id,
        "filename": "RAG_System_Architecture.txt",
        "size": len(SAMPLE_DOCUMENT),
        "uploaded_at": datetime.now().isoformat(),
        "status": "ready",
        "chunk_count": len(chunks2),
        "content": SAMPLE_DOCUMENT
    }
    
    logger.info(f"RAG architecture document loaded: {len(chunks2)} chunks")
    logger.info(f"Total documents: {len(uploaded_documents)}, Total chunks: {len(document_chunks)}")


def load_document_metadata():
    if os.path.exists(METADATA_FILE):
        try:
            with open(METADATA_FILE, 'r') as f:
                return json.load(f)
        except Exception as e:
            logger.warning(f"Failed to load metadata: {e}")
    return {}


def save_document_metadata(docs):
    try:
        os.makedirs(os.path.dirname(METADATA_FILE), exist_ok=True)
        with open(METADATA_FILE, 'w') as f:
            json.dump(docs, f, indent=2)
    except Exception as e:
        logger.error(f"Failed to save metadata: {e}")


# Load existing metadata on startup, THEN load sample documents
uploaded_documents = load_document_metadata()
load_sample_documents()  # Add sample documents after loading metadata


def get_embedding_service():
    global embedding_service
    if embedding_service is None:
        try:
            from sentence_transformers import SentenceTransformer
            embedding_service = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("[OK] Sentence-transformers initialized (384-dim)")
        except Exception as e:
            logger.warning(f"Sentence-transformers failed: {e}")
            embedding_service = "fallback"
    return embedding_service


def get_vector_store():
    global vector_store
    if vector_store is None:
        try:
            from src.embeddings.vector_store_new import VectorStore
            vector_store = VectorStore(db_type="chroma", collection_name="rag_docs")
            logger.info("[OK] Vector store initialized")
        except Exception as e:
            logger.warning(f"Vector store failed: {e}")
            vector_store = "fallback"
    return vector_store


def get_hybrid_retriever():
    global hybrid_retriever
    if hybrid_retriever is None:
        try:
            from src.rag.hybrid_retriever import HybridRetriever
            hybrid_retriever = HybridRetriever(alpha=0.7)
            logger.info("[OK] Hybrid retriever initialized")
        except Exception as e:
            logger.warning(f"Hybrid retriever failed: {e}")
            hybrid_retriever = "fallback"
    return hybrid_retriever


def get_llm_client():
    global llm_client
    if llm_client is None:
        # Try Ollama first (local, free, no API key needed)
        try:
            import ollama
            # Test connection to Ollama
            ollama.list()
            llm_client = ("ollama", ollama)
            logger.info("[OK] Ollama local LLM initialized")
            return llm_client
        except Exception as e:
            logger.info(f"[INFO] Ollama not available: {e}")
        
        # Try OpenAI if Ollama not available
        try:
            from openai import OpenAI
            api_key = os.getenv("OPENAI_API_KEY")
            if api_key:
                llm_client = ("openai", OpenAI(api_key=api_key))
                logger.info("[OK] OpenAI GPT client initialized")
                return llm_client
        except Exception as e:
            logger.warning(f"OpenAI failed: {e}")
        
        # Fallback mode
        logger.info("[INFO] No LLM available - using document excerpts mode")
        llm_client = "fallback"
    return llm_client


@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    x_api_key: Optional[str] = Header(None)
):
    try:
        logger.info(f"Upload request received: {file.filename}")
        
        # Validate file type
        allowed_extensions = ['.pdf', '.docx', '.txt', '.md', '.doc']
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # Read file content
        content = await file.read()
        logger.info(f"File read: {len(content)} bytes")
        
        # === EXTRACT TEXT (FAST - 1-2 seconds max) ===
        text_content, total_pages, extraction_error = extract_text_sync(content, file_ext)
        
        if extraction_error:
            raise HTTPException(status_code=400, detail=f"Failed to extract text: {extraction_error}")
        
        logger.info(f"Text extracted: {len(text_content)} chars from {total_pages} pages")
        
        # Generate document ID
        doc_id = f"doc_{len(uploaded_documents) + 1}_{int(time.time())}"
        file_path = os.path.join(UPLOAD_DIR, f"{doc_id}_{file.filename}")
        
        # Save file to disk
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, 'wb') as f:
            f.write(content)
        
        # === CREATE CHUNKS FOR KEYWORD SEARCH (FAST - no embedding) ===
        chunks = create_smart_chunks(text_content, chunk_size=512, overlap=50)
        
        # Store chunks in memory for keyword search
        for i, chunk_text in enumerate(chunks):
            document_chunks.append({
                "chunk_id": f"{doc_id}_chunk_{i}",
                "doc_id": doc_id,
                "filename": file.filename,
                "content": chunk_text,
                "index": i
            })
        
        # Store document metadata - READY immediately
        uploaded_documents[doc_id] = {
            "id": doc_id,
            "filename": file.filename,
            "file_path": file_path,
            "size": len(content),
            "uploaded_at": datetime.now().isoformat(),
            "status": "ready",  # Ready for keyword search immediately!
            "chunk_count": len(chunks),
            "content": text_content,
            "error": None
        }
        save_document_metadata(uploaded_documents)
        
        logger.info(f"Document {doc_id} ready for search: {len(chunks)} chunks")
        
        # === RETURN SUCCESS - SEARCHABLE NOW ===
        return {
            "status": "ready",
            "doc_id": doc_id,
            "filename": file.filename,
            "message": f"Document ready! You can search it now.",
            "file_size": len(content),
            "pages": total_pages,
            "chunks": len(chunks),
            "text_length": len(text_content),
            "preview": text_content[:300] + "..." if len(text_content) > 300 else text_content
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


def extract_text_sync(content: bytes, file_ext: str) -> tuple:
    try:
        if file_ext == '.pdf':
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text_content = "\n\n".join([page.extract_text() for page in pdf_reader.pages])
            total_pages = len(pdf_reader.pages)
            logger.info(f"PDF: {total_pages} pages extracted")
            
        elif file_ext in ['.docx', '.doc']:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            text_content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            total_pages = len(doc.paragraphs)
            logger.info(f"DOCX: {total_pages} paragraphs extracted")
            
        elif file_ext in ['.txt', '.md']:
            text_content = content.decode('utf-8')
            total_pages = 1
            logger.info(f"Text: {len(text_content)} chars")
        else:
            text_content = content.decode('utf-8', errors='ignore')
            total_pages = 1
            
        return text_content, total_pages, None
    except Exception as e:
        logger.error(f"Text extraction error: {e}")
        return None, 0, str(e)


async def process_document_background(doc_id: str, file_path: str, filename: str, content: bytes, file_ext: str):
    try:
        logger.info(f"Starting background embedding for {doc_id}")
        
        # === RETRIEVE EXTRACTED TEXT ===
        text_content = uploaded_documents[doc_id].get("content")
        if not text_content:
            logger.error(f"No extracted text found for {doc_id}")
            uploaded_documents[doc_id]["status"] = "error"
            uploaded_documents[doc_id]["error"] = "Text content missing"
            save_document_metadata(uploaded_documents)
            return
        
        # === CHUNKING ===
        chunks = create_smart_chunks(text_content, chunk_size=512, overlap=50)
        total_chunks = len(chunks)
        logger.info(f"Created {total_chunks} chunks for embedding")
        
        # === EMBEDDING & INCREMENTAL INDEXING ===
        embedder = get_embedding_service()
        vector_store_instance = get_vector_store()
        
        indexed_count = 0
        batch_size = 10  # Index in batches for better performance
        
        for batch_start in range(0, total_chunks, batch_size):
            batch_end = min(batch_start + batch_size, total_chunks)
            batch_chunks = chunks[batch_start:batch_end]
            
            chunk_ids = []
            chunk_embeddings = []
            chunk_metadata = []
            chunk_texts = []
            
            for i, chunk_text in enumerate(batch_chunks, start=batch_start):
                chunk_id = f"{doc_id}_chunk_{i}"
                
                # Store in memory for keyword search
                document_chunks.append({
                    "chunk_id": chunk_id,
                    "doc_id": doc_id,
                    "filename": filename,
                    "content": chunk_text,
                    "index": i
                })
                
                # Prepare for vector store
                if embedder != "fallback":
                    try:
                        embedding = embedder.encode(chunk_text, convert_to_numpy=True).tolist()
                        chunk_ids.append(chunk_id)
                        chunk_embeddings.append(embedding)
                        chunk_metadata.append({
                            "filename": filename,
                            "doc_id": doc_id,
                            "chunk_index": i
                        })
                        chunk_texts.append(chunk_text)
                    except Exception as e:
                        logger.error(f"Embedding error for chunk {i}: {e}")
            
            # === INCREMENTAL VECTOR STORE INSERTION ===
            if vector_store_instance != "fallback" and chunk_embeddings:
                try:
                    vector_store_instance.add_documents(
                        ids=chunk_ids,
                        embeddings=chunk_embeddings,
                        metadata=chunk_metadata,
                        texts=chunk_texts
                    )
                    indexed_count += len(chunk_embeddings)
                    progress = (batch_end / total_chunks) * 100
                    logger.info(f"Embedded batch {batch_start}-{batch_end}/{total_chunks} ({progress:.1f}%)")
                    
                    # Update progress in metadata
                    uploaded_documents[doc_id]["indexed_chunks"] = indexed_count
                    uploaded_documents[doc_id]["total_chunks"] = total_chunks
                    save_document_metadata(uploaded_documents)
                    
                except Exception as e:
                    logger.warning(f"Vector store batch insertion failed: {e}")
        
        # Update document status to complete
        uploaded_documents[doc_id]["status"] = "completed"
        uploaded_documents[doc_id]["chunk_count"] = total_chunks
        uploaded_documents[doc_id]["indexed_chunks"] = indexed_count
        save_document_metadata(uploaded_documents)
        
        logger.info(f"Background embedding complete for {doc_id}: {indexed_count}/{total_chunks} chunks embedded")
        
    except Exception as e:
        logger.error(f"Background embedding error for {doc_id}: {str(e)}", exc_info=True)
        uploaded_documents[doc_id]["status"] = "error"
        uploaded_documents[doc_id]["error"] = str(e)
        save_document_metadata(uploaded_documents)


def create_smart_chunks(text: str, chunk_size: int = 512, overlap: int = 50) -> List[str]:
    sentences = text.replace('\n', ' ').split('. ')
    chunks = []
    current_chunk = []
    current_size = 0
    
    for sentence in sentences:
        sentence = sentence.strip() + '.'
        sentence_size = len(sentence)
        
        if current_size + sentence_size > chunk_size and current_chunk:
            chunks.append(' '.join(current_chunk))
            # Keep last few sentences for overlap
            overlap_sentences = current_chunk[-2:] if len(current_chunk) > 2 else current_chunk
            current_chunk = overlap_sentences + [sentence]
            current_size = sum(len(s) for s in current_chunk)
        else:
            current_chunk.append(sentence)
            current_size += sentence_size
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks if chunks else [text]


@router.get("/documents/{doc_id}/status")
async def get_document_status(doc_id: str):
    if doc_id not in uploaded_documents:
        raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")
    
    doc = uploaded_documents[doc_id]
    
    # Calculate progress percentage
    progress = 0
    if doc.get("status") in ["processing", "embedding"]:
        total = doc.get("total_chunks", 0)
        indexed = doc.get("indexed_chunks", 0)
        if total > 0:
            progress = (indexed / total) * 100
    elif doc.get("status") == "completed":
        progress = 100
    
    return {
        "doc_id": doc_id,
        "filename": doc.get("filename"),
        "status": doc.get("status"),
        "progress": progress,
        "indexed_chunks": doc.get("indexed_chunks", 0),
        "total_chunks": doc.get("total_chunks", 0),
        "chunk_count": doc.get("chunk_count", 0),
        "uploaded_at": doc.get("uploaded_at"),
        "error": doc.get("error")
    }


@router.get("/documents/{doc_id}/content")
async def get_document_content(doc_id: str, limit: int = None):
    if doc_id not in uploaded_documents:
        raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")
    
    doc = uploaded_documents[doc_id]
    content = doc.get("content", "")
    
    if not content:
        raise HTTPException(status_code=404, detail="Content not yet extracted")
    
    # Return full content or limited by parameter
    if limit:
        content = content[:limit]
    
    return {
        "doc_id": doc_id,
        "filename": doc.get("filename"),
        "content": content,
        "total_length": len(doc.get("content", "")),
        "status": doc.get("status")
    }


@router.get("/documents")
async def list_documents(skip: int = 0, limit: int = 10):
    docs = list(uploaded_documents.values())[skip:skip+limit]
    return {
        "total": len(uploaded_documents),
        "documents": docs
    }


def create_smart_chunks(text: str, chunk_size: int = 512, overlap: int = 50) -> List[str]:
    sentences = text.replace('\n', ' ').split('. ')
    chunks = []
    current_chunk = []
    current_size = 0
    
    for sentence in sentences:
        sentence = sentence.strip() + '.'
        sentence_size = len(sentence)
        
        if current_size + sentence_size > chunk_size and current_chunk:
            chunks.append(' '.join(current_chunk))
            # Keep last few sentences for overlap
            overlap_sentences = current_chunk[-2:] if len(current_chunk) > 2 else current_chunk
            current_chunk = overlap_sentences + [sentence]
            current_size = sum(len(s) for s in current_chunk)
        else:
            current_chunk.append(sentence)
            current_size += sentence_size
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks if chunks else [text]


@router.post("/query", response_model=QueryResponse)
async def query_documents(
    request: QueryRequest,
    x_api_key: Optional[str] = Header(None)
):
    start_time = time.time()
    
    try:
        query = request.query
        
        # Check documents
        if not document_chunks:
            return QueryResponse(
                response="📄 No documents uploaded yet. Upload a PDF/document to ask questions!",
                citations=[],
                confidence_score=0.0,
                retrieved_count=0,
                reranked_count=0,
                processing_time_ms=(time.time() - start_time) * 1000
            )
        
        # Get services
        embedder = get_embedding_service()
        vector_store_instance = get_vector_store()
        hybrid_retriever_instance = get_hybrid_retriever()
        
        # === SEMANTIC SEARCH ===
        semantic_results = []
        if embedder != "fallback":
            try:
                query_embedding = embedder.encode(query, convert_to_numpy=True)
                
                # Use vector store if available
                if vector_store_instance != "fallback":
                    semantic_results = vector_store_instance.search(
                        query_embedding.tolist(), 
                        top_k=request.top_k * 2
                    )
                else:
                    # In-memory semantic search
                    import numpy as np
                    for chunk in document_chunks:
                        if 'embedding' not in chunk:
                            chunk['embedding'] = embedder.encode(chunk['content'], convert_to_numpy=True)
                    
                    scores = []
                    for chunk in document_chunks:
                        similarity = np.dot(query_embedding, chunk['embedding']) / (
                            np.linalg.norm(query_embedding) * np.linalg.norm(chunk['embedding'])
                        )
                        scores.append({
                            'id': chunk['chunk_id'],
                            'score': float(similarity),
                            'text': chunk['content'],
                            'metadata': {
                                'filename': chunk['filename'],
                                'chunk_index': chunk['index']
                            }
                        })
                    
                    scores.sort(key=lambda x: x['score'], reverse=True)
                    semantic_results = scores[:request.top_k * 2]
                
                logger.info(f"Semantic search: {len(semantic_results)} results")
            except Exception as e:
                logger.error(f"Semantic search failed: {e}")
        
        # === HYBRID RETRIEVAL ===
        if hybrid_retriever_instance != "fallback" and semantic_results:
            try:
                final_results = hybrid_retriever_instance.retrieve(
                    query=query,
                    semantic_results=semantic_results,
                    all_chunks=document_chunks,
                    top_k=request.top_k
                )
                logger.info(f"Hybrid retrieval: {len(final_results)} results")
            except Exception as e:
                logger.error(f"Hybrid retrieval failed: {e}")
                final_results = semantic_results[:request.top_k]
        else:
            # Fallback to keyword search when semantic search returns no results
            logger.info("Using keyword search fallback")
            final_results = keyword_search(query, document_chunks)[:request.top_k]
            logger.info(f"Keyword search: {len(final_results)} results")
        
        # === LLM GENERATION ===
        if final_results:
            # Prepare context from top chunks
            context_parts = []
            for i, result in enumerate(final_results[:3], 1):
                text = result.get('text', '')
                filename = result.get('metadata', {}).get('filename', 'Unknown')
                context_parts.append(f"[Source {i} - {filename}]:\n{text}")
            
            context = "\n\n".join(context_parts)
            
            # Get LLM client
            llm = get_llm_client()
            
            if llm != "fallback":
                try:
                    llm_type, llm_instance = llm
                    
                    if llm_type == "ollama":
                        # Generate answer with Ollama (local LLM)
                        # Run in thread so event loop stays free for health checks
                        def _ollama_chat():
                            return llm_instance.chat(
                                model="llama3.2",
                                messages=[
                                    {"role": "system", "content": "You are a helpful AI assistant. Answer questions based ONLY on the provided document context. Always cite which source you used (e.g., [Source 1]). If the answer is not in the context, say so clearly. Be concise and accurate."},
                                    {"role": "user", "content": f"Context from documents:\n\n{context}\n\nQuestion: {query}\n\nProvide a clear, accurate answer based on the context above. Cite sources."}
                                ]
                            )
                        response = await asyncio.to_thread(_ollama_chat)
                        response_text = response['message']['content']
                        logger.info("[OK] Ollama LLM generation complete")
                    
                    elif llm_type == "openai":
                        # Generate answer with OpenAI GPT
                        completion = llm_instance.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": "You are a helpful AI assistant. Answer questions based on the provided document context. Always cite sources. If unsure, say so."},
                                {"role": "user", "content": f"Context:\n\n{context}\n\nQuestion: {query}\n\nProvide a clear, accurate answer based on the context."}
                            ],
                            temperature=0.7,
                            max_tokens=500
                        )
                        response_text = completion.choices[0].message.content
                        logger.info("[OK] OpenAI LLM generation complete")
                    else:
                        response_text = f"**Found relevant information:**\n\n{context[:1500]}"
                        
                except Exception as e:
                    logger.error(f"LLM generation failed: {e}")
                    response_text = f"**Relevant excerpts from your documents:**\n\n{context[:1500]}\n\n*Based on keyword and semantic search from your uploaded documents.*"
            else:
                response_text = f"**Found relevant information:**\n\n{context[:1500]}"
            
            # Build citations
            citations = [
                {
                    "id": result.get('id', ''),
                    "score": result.get('score', 0.0),
                    "metadata": {
                        "filename": result.get('metadata', {}).get('filename', 'Unknown'),
                        "chunk_index": result.get('metadata', {}).get('chunk_index', 0),
                        "preview": result.get('text', '')[:150] + "..."
                    }
                }
                for result in final_results
            ]
            
            confidence = final_results[0].get('score', 0.0) if final_results else 0.0
        else:
            response_text = "❌ No relevant information found. Try:\n• Rephrasing your question\n• Uploading documents with this information\n• Being more specific"
            citations = []
            confidence = 0.0
        
        return QueryResponse(
            response=response_text,
            citations=citations,
            confidence_score=confidence,
            retrieved_count=len(semantic_results) if semantic_results else 0,
            reranked_count=len(final_results),
            processing_time_ms=(time.time() - start_time) * 1000
        )
        
    except Exception as e:
        logger.error(f"Query error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


def keyword_search(query: str, chunks: List[dict]) -> List[dict]:
    query_words = set(query.lower().split())
    scored_chunks = []
    
    for chunk in chunks:
        chunk_words = set(chunk['content'].lower().split())
        overlap = len(query_words & chunk_words)
        if overlap > 0:
            scored_chunks.append({
                'id': chunk.get('chunk_id', ''),
                'score': overlap / len(query_words),
                'text': chunk['content'],
                'metadata': {
                    'filename': chunk.get('filename', 'Unknown'),
                    'chunk_index': chunk.get('index', 0)
                }
            })
    
    scored_chunks.sort(key=lambda x: x['score'], reverse=True)
    return scored_chunks


@router.get("/health", response_model=HealthResponse)
async def health_check():
    return HealthResponse(
        status="healthy",
        version="0.1.0",
        timestamp=datetime.utcnow(),
        services={
            "postgres": "ok",
            "vector_db": "ok",
            "embedding_service": "ok"
        }
    )
