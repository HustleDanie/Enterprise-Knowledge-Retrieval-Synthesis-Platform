import logging
from pathlib import Path
from typing import List, Optional, Dict, Any
from dataclasses import dataclass
from datetime import datetime
import hashlib

logger = logging.getLogger(__name__)


@dataclass
class Document:
    content: str
    metadata: Dict[str, Any]
    doc_id: str
    source: str
    created_at: datetime = None
    
    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.utcnow()


class DocumentLoader:
    
    SUPPORTED_FORMATS = {'.pdf': 'pdf', '.docx': 'docx', '.txt': 'txt', '.md': 'markdown'}
    
    def __init__(self, max_file_size_mb: int = 100):
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
    
    def load_document(self, file_path: str) -> Optional[Document]:
        path = Path(file_path)
        
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        if path.stat().st_size > self.max_file_size_bytes:
            logger.error(f"File too large: {file_path}")
            return None
        
        if path.suffix.lower() not in self.SUPPORTED_FORMATS:
            logger.error(f"Unsupported file format: {path.suffix}")
            return None
        
        try:
            content = self._load_file_content(path)
            doc_id = self._generate_doc_id(file_path, content)
            
            metadata = {
                "file_name": path.name,
                "file_path": str(path),
                "file_size_bytes": path.stat().st_size,
                "file_format": path.suffix.lower(),
                "created_at": datetime.fromtimestamp(path.stat().st_ctime),
                "modified_at": datetime.fromtimestamp(path.stat().st_mtime),
            }
            
            return Document(
                content=content,
                metadata=metadata,
                doc_id=doc_id,
                source=str(path)
            )
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}")
            return None
    
    def load_documents(self, directory: str, recursive: bool = True) -> List[Document]:
        documents = []
        path = Path(directory)
        
        if not path.exists():
            logger.error(f"Directory not found: {directory}")
            return documents
        
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                doc = self.load_document(str(file_path))
                if doc:
                    documents.append(doc)
                    logger.info(f"Loaded document: {file_path.name}")
        
        logger.info(f"Loaded {len(documents)} documents from {directory}")
        return documents
    
    @staticmethod
    def _load_file_content(file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        
        if suffix == '.txt' or suffix == '.md':
            return file_path.read_text(encoding='utf-8')
        elif suffix == '.pdf':
            return DocumentLoader._load_pdf(str(file_path))
        elif suffix == '.docx':
            return DocumentLoader._load_docx(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    @staticmethod
    def _load_pdf(file_path: str) -> str:
        try:
            import PyPDF2
            text = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text())
            return '\n'.join(text)
        except ImportError:
            logger.warning("PyPDF2 not installed, trying pdfplumber")
            try:
                import pdfplumber
                text = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text.append(page.extract_text())
                return '\n'.join(text)
            except ImportError:
                raise ImportError("Please install PyPDF2 or pdfplumber to load PDF files")
    
    @staticmethod
    def _load_docx(file_path: str) -> str:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        except ImportError:
            raise ImportError("Please install python-docx to load DOCX files")
    
    @staticmethod
    def _generate_doc_id(file_path: str, content: str) -> str:
        combined = f"{file_path}{len(content)}"
        return hashlib.md5(combined.encode()).hexdigest()
